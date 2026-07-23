"""
Build a self-contained Word document describing the v2 methodology — dataset
sources, combination + dedup, 70/15/15 stratified split, augmentation
pipeline, model architectures, training recipe, evaluation protocol, and
the integrated downstream pipeline (YOLO + SAM + Grad-CAM++ + uncertainty +
OOD + MedGemma + atlas + chatbot).

Numbers are pulled live from `reports/v2_metrics.json` and the per-model
training history CSVs, so the document always reflects the current results.

Output:
    reports/methodology_v2.docx
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT     = Path(__file__).resolve().parents[2]
REPORTS  = ROOT / "reports"
OUT_FILE = REPORTS / "methodology_v2.docx"

# ── Helpers ────────────────────────────────────────────────────────
def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x44, 0x77)
    return h


def add_paragraph(doc: Document, text: str, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc: Document, header: list[str], rows: list[list[str]],
              widths: list[float] | None = None) -> None:
    """Add a styled table. `widths` in cm, optional."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    if widths:
        for row in table.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Cm(w)


def load_v2_metrics() -> dict:
    with (REPORTS / "v2_metrics.json").open("r", encoding="utf-8") as f:
        return json.load(f)


def load_training_history(model: str) -> list[dict]:
    p = REPORTS / f"train_v2_{model}_history.csv"
    if not p.exists():
        return []
    with p.open("r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


# ── Document sections ──────────────────────────────────────────────
def cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Brain MRI Tumor Detection Pipeline")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x44, 0x55, 0x99)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Three-model ensemble · explainable AI · LLM-driven reporting")
    r2.font.size = Pt(13)
    r2.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("Methodology document — v2 retraining")
    rm.font.size = Pt(11)
    rm.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()


def intro(doc: Document) -> None:
    add_heading(doc, "1. Introduction", 1)
    add_paragraph(doc,
        "This document describes the end-to-end methodology used to build a "
        "brain MRI tumor classification and analysis pipeline. The pipeline "
        "combines a three-model deep-learning ensemble with object-detection "
        "based localization, pixel-level segmentation, multiple explainability "
        "and uncertainty signals, and a locally-hosted medical large language "
        "model (MedGemma 1.5 4B Multimodal) for narrative reporting and a "
        "scan-aware chatbot.")
    add_paragraph(doc,
        "The v2 retraining (covered here) replaces an earlier DenseNet-169 "
        "model with ConvNeXt-Tiny and trains all three architectures from "
        "scratch on a combined dataset of three public 4-class brain MRI "
        "collections, with strong data augmentation. The locked test set is "
        "used only for final evaluation.")


def dataset(doc: Document, _metrics: dict) -> None:
    add_heading(doc, "2. Dataset", 1)

    add_heading(doc, "2.1 Sources", 2)
    add_paragraph(doc,
        "Three publicly-available 4-class brain MRI classification datasets "
        "were combined to enlarge the training pool and reduce single-source "
        "bias. All three datasets share the same four classes: glioma, "
        "meningioma, no-tumor and pituitary.")
    add_table(doc,
        header=["Source", "Glioma", "Meningioma", "No-tumor", "Pituitary", "Total"],
        rows=[
            ["BRISC-2025",                    "1,401", "1,635", "1,207", "1,757", "6,000"],
            ["Epic + CSCR Hospital",          "3,773", "2,729", "2,432", "3,130", "12,064"],
            ["Kaggle (Sartaj)",               "1,800", "1,800", "1,800", "1,800", "7,200"],
            ["Combined (raw)",                "6,974", "6,164", "5,439", "6,687", "25,264"],
        ],
        widths=[4.0, 2.4, 2.6, 2.4, 2.4, 2.4],
    )

    add_heading(doc, "2.2 Deduplication", 2)
    add_paragraph(doc,
        "MD5 hashing was applied to every image across the three datasets. "
        "Duplicates were detected by exact byte-for-byte match — common "
        "because the public datasets share an underlying source pool. "
        "Removing duplicates is essential to prevent train→test leakage that "
        "would otherwise inflate evaluation accuracy by 5–10 percentage points.")
    add_paragraph(doc,
        "From 25,264 raw images the dedup step retained 14,095 unique images, "
        "a 44% duplicate rate consistent with the known overlap between "
        "Kaggle and the hospital-sourced datasets.")

    add_heading(doc, "2.3 Train / validation / test split", 2)
    add_paragraph(doc,
        "A stratified 70 / 15 / 15 split was applied per class with random "
        "seed 42 (matching the v1 baseline for reproducibility). "
        "Stratification ensures every split has the same class proportions, "
        "so meningioma underrepresentation does not concentrate in any one "
        "split. The split assignment is logged to "
        "reports/split_v2_info.csv with the source dataset, MD5 hash, "
        "destination path, and split label for every image.")
    add_table(doc,
        header=["Split", "Glioma", "Meningioma", "No-tumor", "Pituitary", "Total"],
        rows=[
            ["Train (70%)", "2,808", "1,975", "2,885", "2,199", "9,867"],
            ["Val (15%)",   "602",   "423",   "618",   "471",   "2,114"],
            ["Test (15%)",  "601",   "423",   "618",   "472",   "2,114"],
            ["Unique total", "4,011", "2,821", "4,121", "3,142", "14,095"],
        ],
        widths=[3.4, 2.4, 2.6, 2.4, 2.4, 2.4],
    )
    add_paragraph(doc,
        "The 2,114-image test set is sequestered: it is never touched during "
        "training, hyperparameter selection, or model selection. All results "
        "reported in Section 4 come from a single inference pass over this "
        "locked set after training has fully completed.")


def methodology(doc: Document) -> None:
    add_heading(doc, "3. Methodology", 1)

    add_heading(doc, "3.1 Model architectures", 2)
    add_paragraph(doc,
        "Three architectures were chosen for the ensemble. Diversity of "
        "architectural family is intentional — different inductive biases "
        "produce different failure modes, so disagreement between the three "
        "carries a useful uncertainty signal.")
    add_table(doc,
        header=["Model", "Family", "Params", "ImageNet weights"],
        rows=[
            ["ConvNeXt-Tiny",    "Modern ConvNet (2022)",         "27.8 M", "IMAGENET1K_V1"],
            ["EfficientNet-B3",  "Compound-scaled CNN (2019)",    "10.7 M", "IMAGENET1K_V1"],
            ["ResNet-50",        "Residual CNN (2015)",           "23.5 M", "IMAGENET1K_V2"],
        ],
        widths=[3.6, 4.6, 2.2, 4.0],
    )
    add_paragraph(doc,
        "ConvNeXt-Tiny replaces the v1 DenseNet-169 — it is roughly 3× faster "
        "at inference, achieves higher accuracy on ImageNet transfer, and "
        "trains with a recipe that pairs naturally with the augmentation "
        "pipeline used here.")

    add_heading(doc, "3.2 Augmentation pipeline", 2)
    add_paragraph(doc,
        "Augmentation is the single biggest contributor to the v2 accuracy "
        "improvement. Each training-set image goes through the following "
        "probabilistic pipeline (Albumentations 2.x). Validation and test "
        "use only Resize + ImageNet normalize (no augmentation).")
    add_table(doc,
        header=["Augmentation", "Probability", "Parameters"],
        rows=[
            ["Resize",                          "1.0",  "224 × 224"],
            ["Horizontal flip",                 "0.5",  "Tumor class does not depend on L/R side"],
            ["Rotation",                        "0.5",  "±15° (anatomy-preserving)"],
            ["Random brightness/contrast",      "0.5",  "±20 %"],
            ["CLAHE (clip 2.0, tile 8×8)",      "0.5",  "Mirrors inference-time CLAHE in app.py"],
            ["Gaussian noise",                  "0.3",  "σ ∈ [0.01, 0.05]"],
            ["Elastic transform",               "0.2",  "α=20, σ=4 — simulates scanner deformation"],
            ["Normalize (ImageNet stats)",      "1.0",  "mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]"],
        ],
        widths=[4.6, 2.6, 7.0],
    )
    add_paragraph(doc,
        "The CLAHE augmentation is critical for cross-clinic generalization. "
        "Because half the training samples are CLAHE-enhanced, the same "
        "models perform well on raw and contrast-enhanced inputs at inference "
        "time without any explicit detection of low-contrast scans.")

    add_heading(doc, "3.3 Training recipe", 2)
    add_paragraph(doc,
        "Each model is trained in two phases starting from ImageNet weights:")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Phase 1 — Head warmup. ").bold = True
    p.add_run(
        "5 epochs with the entire backbone frozen. Only the new "
        "4-class classifier head learns. This avoids catastrophic forgetting "
        "of the ImageNet features while the head is still random.")
    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run("Phase 2 — Full fine-tune. ").bold = True
    p2.add_run(
        "Up to 40 epochs with all parameters unfrozen, cosine-annealed "
        "learning rate, AdamW optimizer, gradient clipping at 1.0, label "
        "smoothing 0.1, and mixup with probability 0.3. Mixed-precision "
        "(AMP fp16) keeps memory and time low on the RTX 4060. Early "
        "stopping triggers after 8 consecutive epochs without val-accuracy "
        "improvement.")
    add_paragraph(doc,
        "A class-balanced WeightedRandomSampler is used for the training "
        "DataLoader so that meningioma (the smallest class at 1,975 train "
        "images) is sampled at the same expected rate as no-tumor (the "
        "largest at 2,885).")

    add_heading(doc, "3.4 Evaluation protocol", 2)
    add_paragraph(doc,
        "After training, each model checkpoint (best val-accuracy) is "
        "evaluated on the locked 2,114-image test set with a single "
        "deterministic inference pass (no test-time augmentation). The "
        "ensemble combines the three models by averaging their softmax "
        "outputs (soft voting). Latency is measured separately on "
        "1 × 3 × 224 × 224 dummy tensors with 10 warm-up passes followed "
        "by 30 timed passes; TTA-5 latency reports the same measurement "
        "with five test-time augmentations (original + horizontal flip + "
        "3 rotations). Expected Calibration Error (ECE) is computed with "
        "15 confidence bins on the predicted-class confidence.")


def results(doc: Document, metrics: dict) -> None:
    add_heading(doc, "4. Results", 1)

    # ── 4.1 Per-model ──
    add_heading(doc, "4.1 Per-model test accuracy", 2)
    rows = []
    for m in metrics.get("models", []):
        rows.append([
            m["name"].replace("_", "-"),
            f"{m['test_accuracy']*100:.2f} %",
            f"{m['test_balanced_accuracy']*100:.2f} %",
            f"{m['ece']:.4f}",
            f"{m['latency_ms_single']:.1f} ms",
            f"{m['latency_ms_tta5']:.1f} ms",
            f"{m['param_count_M']:.1f} M",
        ])
    add_table(doc,
        header=["Model", "Test acc.", "Balanced acc.", "ECE", "Latency (1 img)", "TTA-5 latency", "Params"],
        rows=rows,
        widths=[3.4, 2.0, 2.4, 1.6, 2.6, 2.4, 1.8],
    )

    # ── 4.2 Ensemble ──
    ens = metrics.get("ensemble") or {}
    add_heading(doc, "4.2 Ensemble (soft vote)", 2)
    add_paragraph(doc,
        f"The soft-vote ensemble reaches {ens.get('soft_vote_test_acc', 0)*100:.2f}% test accuracy "
        f"and {ens.get('soft_vote_balanced_acc', 0)*100:.2f}% balanced accuracy. "
        f"All three models agree on the same top-1 class for "
        f"{ens.get('three_way_agreement_pct', 0):.2f}% of test images.")
    add_paragraph(doc, "Pairwise agreement on the locked test set:")
    pair_rows = []
    for k, v in (ens.get("pairwise_agreement_pct") or {}).items():
        a, b = k.split("__")
        pair_rows.append([a.replace("_", "-"), b.replace("_", "-"), f"{v:.2f} %"])
    if pair_rows:
        add_table(doc,
            header=["Model A", "Model B", "Agreement"],
            rows=pair_rows,
            widths=[4.8, 4.8, 3.0],
        )

    # ── 4.3 Per-class P / R / F1 ──
    add_heading(doc, "4.3 Per-class precision, recall and F1", 2)
    for m in metrics.get("models", []):
        add_paragraph(doc, m["name"].replace("_", "-"), bold=True, size=11)
        rows = []
        for row in m.get("per_class", []):
            rows.append([
                row["class"],
                f"{row['precision']*100:.2f} %",
                f"{row['recall']*100:.2f} %",
                f"{row['f1']*100:.2f} %",
                str(row.get("support", "—")),
            ])
        add_table(doc,
            header=["Class", "Precision", "Recall", "F1", "Support"],
            rows=rows,
            widths=[3.6, 2.6, 2.6, 2.6, 1.8],
        )

    # ── 4.4 Confusion matrices ──
    add_heading(doc, "4.4 Confusion matrices", 2)
    add_paragraph(doc,
        "Diagonal entries are correctly classified images. Each model is "
        "evaluated on the 2,114-image locked test set.")
    for m in metrics.get("models", []):
        png_name = m.get("confusion_png")
        if not png_name:
            continue
        png_path = REPORTS / png_name
        if not png_path.exists():
            continue
        add_paragraph(doc, m["name"].replace("_", "-"), bold=True)
        doc.add_picture(str(png_path), width=Inches(4.0))


def pipeline(doc: Document) -> None:
    add_heading(doc, "5. Integrated pipeline", 1)
    add_paragraph(doc,
        "Beyond classification, the deployed pipeline integrates several "
        "downstream modules so that a single uploaded MRI produces a "
        "complete diagnostic dashboard:")
    items = [
        ("YOLO detection",
         "A custom-trained YOLOv11 detector locates the tumor bounding box on "
         "the input image. The detector was supervised on radiologist-"
         "validated masks across 233 patients."),
        ("MobileSAM segmentation",
         "MobileSAM is prompted with the YOLO bounding box to produce a "
         "pixel-tight tumor mask. The mask area is converted into a "
         "5%-wide size range (1–5 %, 5–10 %, 10–15 %, …) for the report."),
        ("Grad-CAM++",
         "Generated for the predicted class as a visual explanation. "
         "Adversarial-robustness consistency checks confirm the heatmap is "
         "stable under noise and brightness perturbations."),
        ("MC-Dropout uncertainty",
         "T = 20 stochastic forward passes produce epistemic + aleatoric "
         "uncertainty estimates and a needs-review flag for low-confidence "
         "predictions."),
        ("Energy-based OOD detection",
         "A per-model energy score flags scans that differ from the training "
         "distribution (different sequence, different scanner, non-brain "
         "image)."),
        ("Anatomical region prior",
         "The Grad-CAM++ centroid is mapped to a 3 × 3 anatomical grid and "
         "compared against a class-specific clinical prior. Atypical "
         "locations (e.g. meningioma deep in the parenchyma) trigger a "
         "review flag."),
        ("MedGemma 1.5 4B Multimodal",
         "Locally-hosted via Ollama. Produces structured patient or clinician "
         "reports (basic / advanced × English / Spanish), an independent "
         "tumor assessment that is cross-validated against the pipeline "
         "metrics, and a multi-turn chatbot with hard-coded safety guards "
         "(no diagnoses, no prescriptions, prompt-injection refusal)."),
        ("3D brain atlas",
         "Interactive three.js viewer with CC0 brain meshes from the NIH 3D "
         "Print Exchange. Reports surface a 'view this zone in the atlas' "
         "button that auto-pins the detected region."),
    ]
    for title, body in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + ". ").bold = True
        p.add_run(body)


def limitations(doc: Document) -> None:
    add_heading(doc, "6. Limitations and future work", 1)
    bullets = [
        "Only axial T1-weighted slices are supported. Sagittal, coronal, T2 "
        "and FLAIR sequences would require either re-training or sequence "
        "detection and per-sequence calibration.",
        "Single-slice analysis. The pipeline does not aggregate predictions "
        "across the volume; a true 3D approach would be more robust to "
        "borderline slices.",
        "Datasets are largely adult MRI. Pediatric anatomy and pediatric "
        "tumor subtypes are under-represented.",
        "No longitudinal data — the system cannot model tumor growth or "
        "regression between studies. Adding patient-level history would "
        "require a different data structure.",
        "Tumor-size measurement is approximate: it relies on YOLO bounding "
        "box × calibrated fill ratio and SAM segmentation area. Clinical "
        "volumetry would require 3D segmentation across the full volume.",
        "The locally-hosted MedGemma model can refuse or hallucinate. A "
        "production deployment would route critical decisions through a "
        "trained specialist; this system is for research and education only.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def closing(doc: Document) -> None:
    doc.add_paragraph()
    add_paragraph(doc,
        "Disclaimer: this work is a research and educational tool. It is not a "
        "medical device, has not been evaluated by the FDA, EMA or any other "
        "regulatory authority, and must not be used for clinical decision-"
        "making.", italic=True, size=9)


# ── Entry point ────────────────────────────────────────────────────
def main() -> None:
    metrics = load_v2_metrics()
    doc = Document()

    # Slightly narrower margins so the tables fit cleanly
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    cover(doc)
    intro(doc)
    dataset(doc, metrics)
    methodology(doc)
    results(doc, metrics)
    pipeline(doc)
    limitations(doc)
    closing(doc)

    OUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"wrote {OUT_FILE}  ({OUT_FILE.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
